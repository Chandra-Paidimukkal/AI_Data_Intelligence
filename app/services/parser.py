"""
parser.py — ADE-style layout-aware document parsing.

Produces the exact same output schema as LandingAI ADE:
  - markdown: full document as markdown with <a id='...'> anchors
  - chunks:   list of {markdown, type, id, grounding: {box, page}}
  - splits:   page-level splits each with their chunk ids
  - grounding: flat id→{box, page, type, confidence} map
  - metadata:  filename, page_count, duration_ms, etc.

Chunk types (mirroring ADE):
  text | table | figure | logo | title | header | footer | caption

Supports: PDF (PyMuPDF primary, pdfminer fallback), images (OCR), text/md.
"""
from __future__ import annotations

import io
import re
import time
import uuid
from pathlib import Path
from typing import Optional

from loguru import logger

# optional deps
try:
    import fitz
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    from pdfminer.high_level import extract_pages
    from pdfminer.layout import LTTextBox
    PDFMINER_AVAILABLE = True
except ImportError:
    PDFMINER_AVAILABLE = False


# ── public entry-point ──────────────────────────────────────────────────────

def parse_document(file_path: str, mime_type: str = "") -> dict:
    """Main entry point. Returns ADE-compatible parsed document structure."""
    t0 = time.time()
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        result = _parse_pdf(file_path)
    elif suffix in (".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".webp", ".gif", ".heic", ".heif"):
        result = _parse_image(file_path)
    elif suffix in (".txt", ".md", ".rst", ".csv"):
        result = _parse_text(file_path)
    elif suffix in (".docx", ".doc"):
        result = _parse_docx(file_path)
    elif suffix in (".xlsx", ".xls"):
        result = _parse_excel(file_path)
    elif suffix in (".html", ".htm"):
        result = _parse_html(file_path)
    else:
        # Try as text for any unknown format
        result = _parse_text(file_path)

    result["metadata"]["duration_ms"] = int((time.time() - t0) * 1000)
    result["metadata"]["filename"] = path.name
    return result


# ── PDF parser ──────────────────────────────────────────────────────────────

def _parse_pdf(file_path: str) -> dict:
    if PYMUPDF_AVAILABLE:
        return _parse_pdf_pymupdf(file_path)
    elif PDFMINER_AVAILABLE:
        return _parse_pdf_pdfminer(file_path)
    else:
        with open(file_path, "rb") as f:
            text = f.read().decode("utf-8", errors="ignore")
        return _wrap_plain_text(text, Path(file_path).name)


def _parse_pdf_pymupdf(file_path: str) -> dict:
    doc = fitz.open(file_path)
    all_chunks = []
    page_splits = []
    grounding = {}
    ocr_used = False
    page_count = doc.page_count

    for page_num, page in enumerate(doc):
        page_chunks = []
        pw, ph = page.rect.width, page.rect.height

        # text/figure blocks
        raw_dict = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)
        blocks = raw_dict.get("blocks", [])

        for blk in blocks:
            btype = blk.get("type", 0)
            bbox = blk.get("bbox", [0, 0, 0, 0])
            norm = _norm_bbox(bbox, pw, ph)

            if btype == 0:  # text block
                lines_text = []
                for line in blk.get("lines", []):
                    spans_text = " ".join(s.get("text", "") for s in line.get("spans", []))
                    lines_text.append(spans_text)
                raw_text = "\n".join(lines_text).strip()
                if not raw_text:
                    continue

                chunk_type = _classify_text_block(raw_text, blk, page_num)
                chunk_id = _cid()
                md = _text_to_chunk_markdown(raw_text, chunk_type, chunk_id)
                conf = _text_confidence(raw_text)

                chunk = {"markdown": md, "type": chunk_type, "id": chunk_id,
                         "grounding": {"box": norm, "page": page_num}}
                grounding[chunk_id] = {
                    "box": norm, "page": page_num,
                    "type": f"chunk{chunk_type.title()}",
                    "confidence": conf, "low_confidence_spans": []
                }
                page_chunks.append(chunk)

            elif btype == 1:  # image block — OCR it to extract dimension text
                chunk_id = _cid()
                ocr_text = ""
                if OCR_AVAILABLE:
                    try:
                        # Extract the image from the page and OCR it
                        clip = fitz.Rect(bbox)
                        pix = page.get_pixmap(clip=clip, dpi=200)
                        img_data = pix.tobytes("png")
                        img = Image.open(io.BytesIO(img_data))
                        ocr_text = pytesseract.image_to_string(img).strip()
                        # Filter out very short/noisy OCR results
                        if len(ocr_text) < 3:
                            ocr_text = ""
                    except Exception as e:
                        logger.debug(f"Image OCR failed page {page_num}: {e}")

                if ocr_text:
                    # Use OCR text as the chunk content so LandingAI can read dimensions
                    md = f"<a id='{chunk_id}'></a>\n\n[DIAGRAM DIMENSIONS]\n{ocr_text}"
                    chunk_type = "figure"
                    ocr_used = True
                else:
                    img_desc = f"Figure on page {page_num + 1}"
                    md = f"<a id='{chunk_id}'></a>\n\n<::{img_desc}: figure::>"
                    chunk_type = "figure"

                chunk = {"markdown": md, "type": chunk_type, "id": chunk_id,
                         "grounding": {"box": norm, "page": page_num}}
                grounding[chunk_id] = {
                    "box": norm, "page": page_num,
                    "type": "chunkFigure",
                    "confidence": 0.7 if ocr_text else None, "low_confidence_spans": []
                }
                page_chunks.append(chunk)

        # table extraction
        try:
            tab_finder = page.find_tables()
            for tbl in tab_finder.tables:
                rows = tbl.extract()
                if not rows or len(rows) < 2:
                    continue
                tbl_bbox = tbl.bbox
                norm = _norm_bbox(tbl_bbox, pw, ph)
                chunk_id = _cid()
                md_table = _rows_to_markdown_table(rows, chunk_id)
                chunk = {"markdown": md_table, "type": "table", "id": chunk_id,
                         "grounding": {"box": norm, "page": page_num}}
                grounding[chunk_id] = {
                    "box": norm, "page": page_num,
                    "type": "chunkTable",
                    "confidence": 0.95, "low_confidence_spans": []
                }
                # Remove overlapping text chunks
                page_chunks = [c for c in page_chunks
                               if not _bbox_overlap(c["grounding"]["box"], norm, threshold=0.7)]
                page_chunks.append(chunk)
        except Exception as e:
            logger.debug(f"Table extraction page {page_num}: {e}")

        # OCR fallback
        if not page_chunks and OCR_AVAILABLE:
            ocr_text = _ocr_page_pymupdf(page)
            if ocr_text.strip():
                ocr_used = True
                chunk_id = _cid()
                md = f"<a id='{chunk_id}'></a>\n\n{ocr_text.strip()}"
                norm = _norm_bbox([0, 0, pw, ph], pw, ph)
                chunk = {"markdown": md, "type": "text", "id": chunk_id,
                         "grounding": {"box": norm, "page": page_num}}
                grounding[chunk_id] = {
                    "box": norm, "page": page_num,
                    "type": "chunkText",
                    "confidence": 0.7, "low_confidence_spans": []
                }
                page_chunks.append(chunk)

        # sort top→bottom, left→right
        page_chunks.sort(key=lambda c: (
            round(c["grounding"]["box"]["top"], 2),
            round(c["grounding"]["box"]["left"], 2)
        ))

        all_chunks.extend(page_chunks)
        page_splits.append({
            "class": "page",
            "identifier": f"page_{page_num}",
            "pages": [page_num],
            "markdown": "\n\n".join(c["markdown"] for c in page_chunks),
            "chunks": [c["id"] for c in page_chunks]
        })

    doc.close()

    full_markdown = "\n\n<!-- PAGE BREAK -->\n\n".join(s["markdown"] for s in page_splits)
    document_text = _markdown_to_plain(full_markdown)
    # Deduplicate document text (remove repeated paragraphs from multi-column layouts)
    document_text = _deduplicate_text(document_text)
    tables = _chunks_to_tables(all_chunks)
    kv_pairs = _extract_kv_from_text(document_text)
    sections = _detect_sections(document_text)

    return {
        "markdown": full_markdown,
        "chunks": all_chunks,
        "splits": page_splits,
        "grounding": grounding,
        "document_text": document_text,
        "tables": tables,
        "kv_pairs": kv_pairs,
        "sections": sections,
        "layout_blocks": _chunks_to_layout_blocks(all_chunks),
        "pages": [
            {
                "page": i + 1,
                "text": s["markdown"],
                "table_count": sum(1 for cid in s["chunks"]
                                   if grounding.get(cid, {}).get("type") == "chunkTable")
            }
            for i, s in enumerate(page_splits)
        ],
        "metadata": {
            "file_name": Path(file_path).name,
            "page_count": page_count,
            "table_count": sum(1 for c in all_chunks if c["type"] == "table"),
            "ocr_used": ocr_used,
            "version": "ade-2.0",
            "failed_pages": [],
            "credit_usage": page_count * 3,
        }
    }


def _parse_pdf_pdfminer(file_path: str) -> dict:
    all_chunks = []
    page_splits = []
    grounding = {}
    page_count = 0

    for page_num, page_layout in enumerate(extract_pages(file_path)):
        page_count += 1
        pw = float(page_layout.width)
        ph = float(page_layout.height)
        page_chunks = []

        for element in page_layout:
            if isinstance(element, LTTextBox):
                text = element.get_text().strip()
                if not text:
                    continue
                bbox = (element.x0, element.y0, element.x1, element.y1)
                norm = {
                    "left": bbox[0] / pw,
                    "top": 1 - bbox[3] / ph,
                    "right": bbox[2] / pw,
                    "bottom": 1 - bbox[1] / ph,
                }
                chunk_id = _cid()
                ctype = _classify_text_block(text, {}, page_num)
                md = _text_to_chunk_markdown(text, ctype, chunk_id)
                chunk = {"markdown": md, "type": ctype, "id": chunk_id,
                         "grounding": {"box": norm, "page": page_num}}
                grounding[chunk_id] = {
                    "box": norm, "page": page_num,
                    "type": f"chunk{ctype.title()}",
                    "confidence": _text_confidence(text), "low_confidence_spans": []
                }
                page_chunks.append(chunk)

        page_chunks.sort(key=lambda c: (round(c["grounding"]["box"]["top"], 2),
                                        round(c["grounding"]["box"]["left"], 2)))
        all_chunks.extend(page_chunks)
        page_splits.append({
            "class": "page",
            "identifier": f"page_{page_num}",
            "pages": [page_num],
            "markdown": "\n\n".join(c["markdown"] for c in page_chunks),
            "chunks": [c["id"] for c in page_chunks]
        })

    full_markdown = "\n\n<!-- PAGE BREAK -->\n\n".join(s["markdown"] for s in page_splits)
    document_text = _markdown_to_plain(full_markdown)

    return {
        "markdown": full_markdown,
        "chunks": all_chunks,
        "splits": page_splits,
        "grounding": grounding,
        "document_text": document_text,
        "tables": _chunks_to_tables(all_chunks),
        "kv_pairs": _extract_kv_from_text(document_text),
        "sections": _detect_sections(document_text),
        "layout_blocks": _chunks_to_layout_blocks(all_chunks),
        "pages": [{"page": i + 1, "text": s["markdown"], "table_count": 0}
                  for i, s in enumerate(page_splits)],
        "metadata": {
            "file_name": Path(file_path).name,
            "page_count": page_count,
            "table_count": 0,
            "ocr_used": False,
            "version": "ade-2.0",
            "failed_pages": [],
            "credit_usage": page_count * 3,
        }
    }


# ── Image / text parsers ────────────────────────────────────────────────────

def _parse_image(file_path: str) -> dict:
    text = ""
    ocr_used = False
    chunk_id = _cid()

    if OCR_AVAILABLE:
        img = Image.open(file_path)
        text = pytesseract.image_to_string(img)
        ocr_used = True

    norm = {"left": 0.0, "top": 0.0, "right": 1.0, "bottom": 1.0}
    md = _text_to_chunk_markdown(text.strip(), "text", chunk_id) if text.strip() else ""
    chunk = {"markdown": md, "type": "text", "id": chunk_id,
             "grounding": {"box": norm, "page": 0}}
    grounding = {chunk_id: {"box": norm, "page": 0, "type": "chunkText",
                             "confidence": 0.8 if ocr_used else None,
                             "low_confidence_spans": []}}
    split = {"class": "page", "identifier": "page_0", "pages": [0],
             "markdown": md, "chunks": [chunk_id]}

    return {
        "markdown": md,
        "chunks": [chunk],
        "splits": [split],
        "grounding": grounding,
        "document_text": text,
        "tables": _extract_tables_from_text(text),
        "kv_pairs": _extract_kv_from_text(text),
        "sections": _detect_sections(text),
        "layout_blocks": [],
        "pages": [{"page": 1, "text": text, "table_count": 0}],
        "metadata": {
            "file_name": Path(file_path).name,
            "page_count": 1,
            "table_count": 0,
            "ocr_used": ocr_used,
            "version": "ade-2.0",
            "failed_pages": [],
            "credit_usage": 3,
        }
    }


def _parse_text(file_path: str) -> dict:
    with open(file_path, "r", errors="ignore") as f:
        text = f.read()
    return _wrap_plain_text(text, Path(file_path).name)


def _wrap_plain_text(text: str, filename: str) -> dict:
    chunk_id = _cid()
    norm = {"left": 0.0, "top": 0.0, "right": 1.0, "bottom": 1.0}
    md = _text_to_chunk_markdown(text.strip(), "text", chunk_id)
    chunk = {"markdown": md, "type": "text", "id": chunk_id,
             "grounding": {"box": norm, "page": 0}}
    grounding = {chunk_id: {"box": norm, "page": 0, "type": "chunkText",
                             "confidence": 1.0, "low_confidence_spans": []}}
    split = {"class": "page", "identifier": "page_0", "pages": [0],
             "markdown": md, "chunks": [chunk_id]}
    tables = _extract_tables_from_text(text)
    return {
        "markdown": md,
        "chunks": [chunk],
        "splits": [split],
        "grounding": grounding,
        "document_text": text,
        "tables": tables,
        "kv_pairs": _extract_kv_from_text(text),
        "sections": _detect_sections(text),
        "layout_blocks": [],
        "pages": [{"page": 1, "text": text, "table_count": len(tables)}],
        "metadata": {
            "file_name": filename,
            "page_count": 1,
            "table_count": len(tables),
            "ocr_used": False,
            "version": "ade-2.0",
            "failed_pages": [],
            "credit_usage": 3,
        }
    }


# ── Classification ──────────────────────────────────────────────────────────

_TITLE_RE = re.compile(r"^[A-Z][A-Z\s\-/&]{4,80}$")


def _classify_text_block(text: str, blk: dict, page_num: int) -> str:
    stripped = text.strip()
    if len(stripped) < 80 and _TITLE_RE.match(stripped):
        return "title"
    if "\n" not in stripped and ("®" in stripped or "™" in stripped) and len(stripped) < 60:
        return "logo"
    return "text"


def _text_confidence(text: str) -> float:
    if not text.strip():
        return 0.0
    garbage = len(re.findall(r"[^\x20-\x7E\u00A0-\uFFFF]", text))
    ratio = garbage / max(len(text), 1)
    return round(max(0.5, 1.0 - ratio * 5), 3)


# ── Markdown helpers ────────────────────────────────────────────────────────

def _text_to_chunk_markdown(text: str, chunk_type: str, chunk_id: str) -> str:
    anchor = f"<a id='{chunk_id}'></a>"
    if chunk_type == "title":
        return f"{anchor}\n\n## {text.strip()}"
    elif chunk_type == "logo":
        escaped = text.strip().replace("\n", " ")
        return f"{anchor}\n\n<::logo: {escaped}::>"
    else:
        return f"{anchor}\n\n{text.strip()}"


def _rows_to_markdown_table(rows: list, chunk_id: str) -> str:
    anchor = f"<a id='{chunk_id}'></a>"
    tbl_id = chunk_id[:8]
    lines = [f'<table id="{tbl_id}">']
    for r_idx, row in enumerate(rows):
        lines.append("<tr>")
        tag = "th" if r_idx == 0 else "td"
        for c_idx, cell in enumerate(row):
            cell_id = f"{tbl_id}-{r_idx}-{c_idx}"
            val = str(cell).strip() if cell is not None else ""
            lines.append(f'  <{tag} id="{cell_id}">{val}</{tag}>')
        lines.append("</tr>")
    lines.append("</table>")
    return f"{anchor}\n\n" + "\n".join(lines)


def _markdown_to_plain(md: str) -> str:
    text = re.sub(r"<a id='[^']*'></a>", "", md)
    text = re.sub(r"<[^>]+>", " ", text)
    text = re.sub(r"<::.*?::>", "", text, flags=re.DOTALL)
    text = re.sub(r"<!-- PAGE BREAK -->", "\n\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ── BBox helpers ────────────────────────────────────────────────────────────

def _norm_bbox(bbox, pw: float, ph: float) -> dict:
    x0, y0, x1, y1 = bbox[:4]
    return {
        "left": max(0.0, min(1.0, x0 / pw)),
        "top": max(0.0, min(1.0, y0 / ph)),
        "right": max(0.0, min(1.0, x1 / pw)),
        "bottom": max(0.0, min(1.0, y1 / ph)),
    }


def _bbox_overlap(a: dict, b: dict, threshold: float = 0.5) -> bool:
    ox = max(0, min(a["right"], b["right"]) - max(a["left"], b["left"]))
    oy = max(0, min(a["bottom"], b["bottom"]) - max(a["top"], b["top"]))
    intersection = ox * oy
    area_a = (a["right"] - a["left"]) * (a["bottom"] - a["top"])
    if area_a <= 0:
        return False
    return (intersection / area_a) >= threshold


# ── Legacy pipeline converters ──────────────────────────────────────────────

def _chunks_to_tables(chunks: list) -> list:
    tables = []
    for chunk in chunks:
        if chunk["type"] != "table":
            continue
        md = chunk["markdown"]
        rows_raw = re.findall(r"<tr>(.*?)</tr>", md, re.DOTALL)
        parsed_rows = []
        for row in rows_raw:
            cells = re.findall(r"<t[hd][^>]*>(.*?)</t[hd]>", row, re.DOTALL)
            parsed_rows.append([re.sub(r"<[^>]+>", "", c).strip() for c in cells])
        if len(parsed_rows) >= 2:
            headers = parsed_rows[0]
            if _is_layout_table(headers, parsed_rows):
                continue
            # Normalize headers: collapse newlines, strip whitespace
            clean_headers = [_normalize_header(h) for h in headers]
            # Normalize data cells
            clean_rows = [[_normalize_cell(c) for c in row] for row in parsed_rows[1:]]
            tables.append({
                "headers": clean_headers,
                "rows": clean_rows,
                "raw": parsed_rows,
                "chunk_id": chunk["id"],
                "page": chunk["grounding"]["page"]
            })
    return tables


def _normalize_header(h: str) -> str:
    """Normalize a table header: collapse newlines, strip extra whitespace."""
    if not h:
        return h
    # Replace newlines with space
    h = re.sub(r"\s*\n\s*", " ", h).strip()
    # Remove trailing punctuation noise
    h = h.strip(".")
    return h


def _normalize_cell(cell: str) -> str:
    """
    Normalize a table cell value.
    - Keep the primary value before the metric equivalent in parentheses
      e.g. '24"\n(610mm)' → '24"'  but '265 lbs.\n(120kg)' → '265 lbs.'
    - Collapse internal newlines to space for multi-line cells
    """
    if not cell:
        return cell
    # If cell has a newline followed by a parenthesized metric, keep first line
    m = re.match(r'^(.+?)\n\([\d\.]+\s*(?:mm|cm|kg|lbs?)\)', cell.strip())
    if m:
        return m.group(1).strip()
    # Otherwise collapse newlines to space
    return re.sub(r"\s*\n\s*", " ", cell).strip()


def _is_layout_table(headers: list, rows: list) -> bool:
    """
    Detect PDF layout artifacts masquerading as tables.
    A layout table typically has 1-2 columns where one cell contains
    a large block of text (spec content, bullet lists, etc.).
    """
    col_count = len(headers)
    # Only flag 1-2 column tables
    if col_count > 3:
        return False
    # Check if any cell across all rows is a large blob (>300 chars)
    for row in rows:
        for cell in row:
            if len(str(cell).strip()) > 300:
                return True
    return False


def _chunks_to_layout_blocks(chunks: list) -> list:
    return [
        {
            "page": c["grounding"]["page"] + 1,
            "type": c["type"],
            "bbox": [
                c["grounding"]["box"]["left"],
                c["grounding"]["box"]["top"],
                c["grounding"]["box"]["right"],
                c["grounding"]["box"]["bottom"],
            ],
            "text": _markdown_to_plain(c["markdown"])[:200],
            "chunk_id": c["id"],
        }
        for c in chunks
    ]


def _extract_tables_from_text(text: str) -> list:
    tables = []
    lines = text.split("\n")
    current_table = []
    for line in lines:
        if "|" in line and line.count("|") >= 2:
            current_table.append(line)
        else:
            if len(current_table) >= 2:
                t = _parse_pipe_table(current_table)
                if t:
                    tables.append(t)
            current_table = []
    if len(current_table) >= 2:
        t = _parse_pipe_table(current_table)
        if t:
            tables.append(t)
    return tables


def _parse_pipe_table(lines: list) -> Optional[dict]:
    try:
        rows = []
        for line in lines:
            if re.match(r"^[\|\-\+\s]+$", line):
                continue
            cells = [c.strip() for c in line.strip("|").split("|")]
            rows.append(cells)
        if not rows:
            return None
        return {"headers": rows[0], "rows": rows[1:], "raw": lines}
    except Exception:
        return None


def _extract_kv_from_text(text: str) -> list:
    """
    Extract key-value pairs from document text.
    Handles patterns like:
      - 'Key: Value'
      - 'Key = Value'
      - 'Product Name: Master Electric Griddle'
    Filters out garbage multi-line keys and very short/long keys.
    """
    kv = []
    seen = set()

    # Pattern 1: standard "Key: Value" on same line
    for m in re.finditer(
        r"^([A-Za-z][A-Za-z0-9\s\-_/\.]{1,50}?)\s*[:]\s*([^\n]{1,300})",
        text, re.MULTILINE
    ):
        key = m.group(1).strip()
        value = m.group(2).strip()
        # Skip keys with newlines (multi-line garbage)
        if "\n" in key:
            continue
        # Skip very short keys (single word noise) or pure numbers
        if len(key) < 2 or key.isdigit():
            continue
        # Skip if key looks like a sentence fragment
        if key.count(" ") > 5:
            continue
        key_lower = key.lower()
        if key_lower not in seen and value:
            kv.append({"key": key, "value": value})
            seen.add(key_lower)

    # Pattern 2: "Label = Value"
    for m in re.finditer(
        r"([A-Za-z][A-Za-z0-9\s\-_]{1,40})\s*=\s*([^\n]{1,200})",
        text
    ):
        key = m.group(1).strip()
        value = m.group(2).strip()
        if "\n" in key or len(key) < 2:
            continue
        key_lower = key.lower()
        if key_lower not in seen and value:
            kv.append({"key": key, "value": value})
            seen.add(key_lower)

    return kv


def _deduplicate_text(text: str) -> str:
    """
    Remove duplicate paragraphs that appear from multi-column PDF layouts
    where PyMuPDF reads the same content twice (once per column).
    Keeps first occurrence of any paragraph seen more than once.
    """
    paragraphs = re.split(r"\n{2,}", text)
    seen: set = set()
    unique: list = []
    for para in paragraphs:
        # Normalize for comparison: lowercase, collapse whitespace
        key = re.sub(r"\s+", " ", para.strip().lower())
        if len(key) < 20:
            # Keep short paragraphs always (headers, labels)
            unique.append(para)
            continue
        if key not in seen:
            seen.add(key)
            unique.append(para)
    return "\n\n".join(unique)


def _detect_sections(text: str) -> list:
    """
    Detect section headings in document text.
    Catches both ALL-CAPS headings and Title Case headings ending with ':'.
    """
    sections = []
    seen_starts = set()

    # Pattern 1: ALL CAPS headings (e.g. "STANDARD FEATURES", "CLEARANCES")
    for m in re.finditer(r"^([A-Z][A-Z\s\-/]{3,60})\s*$", text, re.MULTILINE):
        title = m.group(1).strip()
        if m.start() not in seen_starts:
            sections.append({"title": title, "start": m.start(), "end": m.end()})
            seen_starts.add(m.start())

    # Pattern 2: Title Case headings ending with ':' (e.g. "Standard Features:", "Specifications:")
    for m in re.finditer(r"^([A-Z][A-Za-z\s]{3,50}):\s*$", text, re.MULTILINE):
        title = m.group(1).strip()
        if m.start() not in seen_starts:
            sections.append({"title": title, "start": m.start(), "end": m.end()})
            seen_starts.add(m.start())

    # Sort by position
    sections.sort(key=lambda s: s["start"])
    return sections


def _ocr_page_pymupdf(page) -> str:
    try:
        pix = page.get_pixmap(dpi=200)
        img_data = pix.tobytes("png")
        img = Image.open(io.BytesIO(img_data))
        return pytesseract.image_to_string(img)
    except Exception as e:
        logger.debug(f"OCR failed: {e}")
        return ""


def _cid() -> str:
    return str(uuid.uuid4())


# ── DOCX parser ─────────────────────────────────────────────────────────────

def _parse_docx(file_path: str) -> dict:
    """Parse Word documents (.docx)."""
    try:
        import docx
        doc = docx.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables_text = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            if rows:
                tables_text.append(rows)
        text = "\n\n".join(paragraphs)
        # Add table content to text
        for t in tables_text:
            for row in t:
                text += "\n" + " | ".join(row)
        return _wrap_plain_text(text, Path(file_path).name)
    except ImportError:
        logger.warning("python-docx not installed, falling back to text parse")
        return _parse_text(file_path)
    except Exception as e:
        logger.error(f"DOCX parse error: {e}")
        return _wrap_plain_text("", Path(file_path).name)


# ── Excel parser ─────────────────────────────────────────────────────────────

def _parse_excel(file_path: str) -> dict:
    """Parse Excel files (.xlsx, .xls)."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        all_text = []
        all_tables = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                clean = [str(c).strip() if c is not None else "" for c in row]
                if any(c for c in clean):
                    rows.append(clean)
            if rows:
                all_text.append(f"Sheet: {sheet_name}")
                for row in rows:
                    all_text.append(" | ".join(row))
                if len(rows) >= 2:
                    all_tables.append({
                        "headers": rows[0],
                        "rows": rows[1:],
                        "raw": rows,
                        "chunk_id": _cid(),
                        "page": 0,
                    })
        text = "\n".join(all_text)
        result = _wrap_plain_text(text, Path(file_path).name)
        result["tables"] = all_tables
        result["metadata"]["table_count"] = len(all_tables)
        return result
    except ImportError:
        logger.warning("openpyxl not installed")
        return _wrap_plain_text("", Path(file_path).name)
    except Exception as e:
        logger.error(f"Excel parse error: {e}")
        return _wrap_plain_text("", Path(file_path).name)


# ── HTML parser ──────────────────────────────────────────────────────────────

def _parse_html(file_path: str) -> dict:
    """Parse HTML files."""
    try:
        with open(file_path, "r", errors="ignore") as f:
            html = f.read()
        # Strip tags
        import re as _re
        text = _re.sub(r"<script[^>]*>.*?</script>", "", html, flags=_re.DOTALL)
        text = _re.sub(r"<style[^>]*>.*?</style>", "", text, flags=_re.DOTALL)
        text = _re.sub(r"<[^>]+>", " ", text)
        text = _re.sub(r"\s{2,}", " ", text).strip()
        return _wrap_plain_text(text, Path(file_path).name)
    except Exception as e:
        logger.error(f"HTML parse error: {e}")
        return _wrap_plain_text("", Path(file_path).name)
